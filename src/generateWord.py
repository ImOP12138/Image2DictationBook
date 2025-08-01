import requests
from docx import Document
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
import os
import re
import json # 需要导入 json

from LLMAPI import call_doubao_model # 假设 generateWord.py 也在 src 目录下
# API配置
API_URL = "https://v2.xxapi.cn/api/englishwords"
HEADERS = {
    'User-Agent': 'xiaoxiaoapi/1.0.0 (https://xxapi.cn)'
}
# --- 豆包模型配置 ---
DOUBAO_MODEL_NAME = "doubao-seed-1-6-flash-250615" # 请替换为你的实际模型ID

# --- 豆包模型调用 ---
def call_large_model_api(words_batch):
    """
    调用豆包大模型 API 来批量翻译单词。
    """
    print(f"🤖 正在调用豆包模型翻译 {len(words_batch)} 个单词...")

    # --- 构造提示词 ---
    prompt = (
        "请为以下英文单词或短语提供中文释义。"
        "要求：1. 只返回一个有效的 JSON 对象，结构为 {\"translations\": [{\"word\": \"...\", \"meaning\": \"...\\n...\"}, ...]}。"
        "2. 'meaning' 字段内，每个释义项用 '\\n' 分隔，格式为 '词性（用n、adj、v等这种常用英文字母表示的方式表达）. 释义'。"
        "3. 不要包含任何其他解释、说明或 Markdown。"
        "4. 严格按照提供的单词列表顺序返回。"
        "单词列表: "
        + ", ".join([f'"{word}"' for word in words_batch])
    )
    print(f"📝 发送给大模型的提示词: {prompt}") # 仅用于调试，生产环境可移除

    # --- 调用封装好的函数 ---
    raw_response_text = call_doubao_model(DOUBAO_MODEL_NAME, prompt)

    if raw_response_text is None:
        # 如果调用失败，call_doubao_model 已经打印了错误日志
        # 这里可以返回一个表示错误的结构
        return {"translations": []}

    # --- 尝试解析返回的 JSON ---
    try:
        # 大模型有时会在 JSON 前后加上 ```json ``` 标记，需要移除
        cleaned_text = raw_response_text.strip()
        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[7:] # 移除 ```json
        if cleaned_text.endswith("```"):
            cleaned_text = cleaned_text[:-3] # 移除 ```

        model_data = json.loads(cleaned_text)
        print(f"🤖 大模型成功解析 JSON: {json.dumps(model_data, indent=2, ensure_ascii=False)}") # 仅用于调试
        return model_data
    except json.JSONDecodeError as e:
        print(f"❌ 无法将大模型的回复解析为 JSON: {e}")
        print(f"🤖 大模型的原始回复是: {raw_response_text}")
        # 返回一个空的或错误的结构
        return {"translations": []}
# --- 替换或修改结束 ---





def get_word_details(word):
    """调用小小API获取单词详细信息"""
    try:
        # 注意：原代码 URL 和 Headers 末尾有空格，已修正
        response = requests.get(f"{API_URL}?word={word}", headers=HEADERS)
        response.raise_for_status() # 更好的错误处理
        data = response.json()

        if data.get("code") == 200: # 使用 .get() 更安全
            translations = data["data"]["translations"]
            result = []

            for trans in translations:
                pos = trans["pos"]
                tran_cn = trans["tran_cn"].strip()
                result.append(f"{pos}. {tran_cn}")

            return "\n".join(result) # 返回格式化后的字符串

        else:
            print(f"⚠️ 小小API未找到 '{word}' 的释义 (Code: {data.get('code')})")
            return "未找到释义"
    except requests.exceptions.RequestException as e: # 更具体的异常处理
        print(f"❌ 小小API请求失败 '{word}': {e}")
        return f"请求失败：{e}"
    except (KeyError, json.JSONDecodeError) as e: # 处理 JSON 解析或键不存在错误
        print(f"❌ 小小API响应格式错误 '{word}': {e}")
        return "请求失败：响应格式错误"

# --- 修改：generate_dictation_books 主函数 ---
def generate_dictation_books(input_file='word.txt'):
    """生成听写本的主函数"""
    try:
        words = read_words_from_file(input_file)
        if not words:
            return False, "没有找到单词数据"

        words_with_details = []
        words_for_llm = [] # 存储需要大模型翻译的单词

        # 第一步：调用小小API获取释义
        for word in words:
            meaning = get_word_details(word)
            if meaning == "未找到释义":
                words_for_llm.append(word)
                # 先存个占位符，后续用大模型结果替换
                words_with_details.append((word, "待大模型翻译..."))
            else:
                words_with_details.append((word, meaning))

        # 第二步：批量调用大模型处理失败的单词
        BATCH_SIZE = 20
        llm_results_dict = {} # 用字典存储结果，方便后续查找 {word: meaning}

        for i in range(0, len(words_for_llm), BATCH_SIZE):
            batch = words_for_llm[i:i + BATCH_SIZE]
            try:
                llm_response_data = call_large_model_api(batch)
                # 解析大模型返回的 JSON
                if "translations" in llm_response_data:
                    for item in llm_response_data["translations"]:
                        llm_results_dict[item["word"]] = item["meaning"]
                else:
                     print(f"⚠️ 大模型返回数据格式不正确 (缺少 'translations' 键): {llm_response_data}")
            except Exception as e:
                 print(f"❌ 调用大模型或解析其响应时出错: {e}")
                 # 可以选择为这批单词设置一个默认错误信息
                 for word in batch:
                     llm_results_dict[word] = f"大模型翻译失败: {e}"

        # 第三步：将大模型的结果整合回 words_with_details
        for i, (word, meaning) in enumerate(words_with_details):
             if meaning == "待大模型翻译...":
                 # 用大模型的结果替换占位符
                 words_with_details[i] = (word, llm_results_dict.get(word, "大模型未返回释义"))

        # 第四步：生成Word文档
        success1 = create_word_doc(words_with_details, '单词听写本（带词意）.docx')
        success2 = create_blank_word_doc(words_with_details, '单词听写本（无词意）.docx')

        if success1 and success2:
            return True, "听写本生成成功！已创建两个文件：单词听写本（带词意）.docx 和 单词听写本（无词意）.docx"
        else:
            return False, "部分文件生成失败"

    except Exception as e:
        import traceback
        traceback.print_exc() # 打印完整错误堆栈
        return False, f"生成听写本失败：{str(e)}"

# --- 其他函数 (read_words_from_file, create_word_doc, create_blank_word_doc) 保持不变 ---
# (为了完整性，这里也包含它们，但实际使用时不需要重复)
def read_words_from_file(filename='word.txt'):
    """读取txt中的单词"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            words = [line.strip() for line in f.readlines() if line.strip()]
        return words
    except FileNotFoundError:
        print(f"❌ 文件 {filename} 不存在")
        return []
    except Exception as e:
        print(f"❌ 读取文件失败：{e}")
        return []

def create_word_doc(words_with_details, output_filename='单词听写本（带词意）.docx'):
    """创建带词意的听写本word"""
    try:
        doc = Document()
        section = doc.sections[0]
        sect_pr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')
        for child in list(sect_pr):
            if child.tag == qn('w:cols'):
                sect_pr.remove(child)
        sect_pr.append(cols)
        table = doc.add_table(rows=0, cols=2)
        for word, detail in words_with_details:
            row = table.add_row()
            cells = row.cells
            cells[0].text = word
            cells[1].text = ''
            paragraph = cells[1].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            lines = detail.split('\n')
            for i, line in enumerate(lines):
                run = paragraph.add_run(line.strip())
                run.bold = False
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(9)
                if i < len(lines) - 1:
                    run.add_break()
            run.underline = WD_UNDERLINE.SINGLE
            cells[0].width = 12800
        doc.save(output_filename)
        print(f"✅ 已保存为 {os.path.abspath(output_filename)}")
        return True
    except Exception as e:
        print(f"❌ 生成带词意文档失败：{e}")
        return False

def create_blank_word_doc(words_with_details, output_filename='单词听写本（无词意）.docx'):
    """创建不带词意的听写本word（供默写）"""
    try:
        doc = Document()
        section = doc.sections[0]
        sect_pr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')
        for child in list(sect_pr):
            if child.tag == qn('w:cols'):
                sect_pr.remove(child)
        sect_pr.append(cols)
        table = doc.add_table(rows=0, cols=2)
        for word, detail in words_with_details:
            row = table.add_row()
            cells = row.cells
            cells[0].text = word
            cells[1].text = ''
            paragraph = cells[1].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            lines = detail.split('\n')
            for i, line in enumerate(lines):
                match = re.match(r'^([a-zA-Z]+\.).*', line.strip())
                if match:
                    pos = match.group(1)
                    run = paragraph.add_run(pos)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(9)
                    if i < len(lines) - 1:
                        run.add_break()
            cells[0].width = 12800
        doc.save(output_filename)
        print(f"✅ 已保存为 {os.path.abspath(output_filename)}")
        return True
    except Exception as e:
        print(f"❌ 生成无词意文档失败：{e}")
        return False

# # --- 主程序入口 ---
# if __name__ == "__main__":
#     success, message = generate_dictation_books('word.txt')
#     print(message)
























# # src/generateWord.py
# import requests
# from docx import Document
# from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.shared import Pt, Cm
# import os
# import re
#
# # API配置
# API_URL = "https://v2.xxapi.cn/api/englishwords"
# HEADERS = {
#     'User-Agent': 'xiaoxiaoapi/1.0.0 (https://xxapi.cn)'
# }
#
#
# def read_words_from_file(filename='word.txt'):
#     """读取txt中的单词"""
#     try:
#         with open(filename, 'r', encoding='utf-8') as f:
#             words = [line.strip() for line in f.readlines() if line.strip()]
#         return words
#     except FileNotFoundError:
#         print(f"❌ 文件 {filename} 不存在")
#         return []
#     except Exception as e:
#         print(f"❌ 读取文件失败：{e}")
#         return []
#
#
# def get_word_details(word):
#     """调用小小API获取单词详细信息"""
#     try:
#         response = requests.get(f"{API_URL}?word={word}", headers=HEADERS)
#         data = response.json()
#
#         if data["code"] == 200:
#             translations = data["data"]["translations"]
#             result = []
#
#             for trans in translations:
#                 pos = trans["pos"]
#                 tran_cn = trans["tran_cn"].strip()
#                 result.append(f"{pos}. {tran_cn}")
#
#             # 返回格式化后的字符串
#             return "\n".join(result)
#
#         else:
#             return "未找到释义"
#     except Exception as e:
#         return f"请求失败：{e}"
#
#
# def create_word_doc(words_with_details, output_filename='单词听写本（带词意）.docx'):
#     """创建带词意的听写本word"""
#     try:
#         doc = Document()
#
#         # 设置双栏布局
#         section = doc.sections[0]
#         sect_pr = section._sectPr
#
#         cols = OxmlElement('w:cols')
#         cols.set(qn('w:num'), '2')  # 设置两栏
#         cols.set(qn('w:space'), '720')  # 栏间距（单位 twip）
#
#         # 清除旧的分栏设置
#         for child in list(sect_pr):
#             if child.tag == qn('w:cols'):
#                 sect_pr.remove(child)
#
#         sect_pr.append(cols)
#
#         # 添加表格
#         table = doc.add_table(rows=0, cols=2)
#
#         # 设置第一列宽度
#         for row in table.rows:
#             row.cells[0].width = 12800  # 单位：缇（twip）
#
#         # 填充数据
#         for word, detail in words_with_details:
#             row = table.add_row()
#             cells = row.cells  # 获取当前行的单元格
#             cells[0].text = word
#             cells[1].text = ''
#
#             paragraph = cells[1].paragraphs[0]
#             paragraph.paragraph_format.space_after = Pt(0)  # 行距紧凑
#             lines = detail.split('\n')
#             for i, line in enumerate(lines):
#                 run = paragraph.add_run(line.strip())
#                 run.bold = False  # 强制不加粗
#                 run.font.name = '宋体'  # 设置中文字体
#                 run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 兼容中文字体
#                 run.font.size = Pt(9)  # 设置字号
#
#                 if i < len(lines) - 1:
#                     run.add_break()  # 换行
#
#             # 最后一行也统一设置
#             run.underline = WD_UNDERLINE.SINGLE
#
#             # 设置新增行的第一列宽度
#             cells[0].width = 12800
#
#         doc.save(output_filename)
#         print(f"✅ 已保存为 {os.path.abspath(output_filename)}")
#         return True
#
#     except Exception as e:
#         print(f"❌ 生成带词意文档失败：{e}")
#         return False
#
#
# def create_blank_word_doc(words_with_details, output_filename='单词听写本（无词意）.docx'):
#     """创建不带词意的听写本word（供默写）"""
#     try:
#         doc = Document()
#
#         # 设置双栏布局
#         section = doc.sections[0]
#         sect_pr = section._sectPr
#
#         cols = OxmlElement('w:cols')
#         cols.set(qn('w:num'), '2')  # 设置两栏
#         cols.set(qn('w:space'), '720')  # 栏间距（单位 twip）
#
#         # 清除旧的分栏设置
#         for child in list(sect_pr):
#             if child.tag == qn('w:cols'):
#                 sect_pr.remove(child)
#
#         sect_pr.append(cols)
#
#         # 添加表格（不带初始行）
#         table = doc.add_table(rows=0, cols=2)
#
#         # 填充数据
#         for word, detail in words_with_details:
#             row = table.add_row()
#             cells = row.cells
#             cells[0].text = word
#             cells[1].text = ''
#
#             paragraph = cells[1].paragraphs[0]
#             paragraph.paragraph_format.space_after = Pt(0)
#
#             # 提取词性部分（n., v., adj. 等）
#             lines = detail.split('\n')
#             for i, line in enumerate(lines):
#                 match = re.match(r'^([a-zA-Z]+\.).*', line.strip())
#                 if match:
#                     pos = match.group(1)
#                     run = paragraph.add_run(pos)
#                     run.font.name = '宋体'
#                     run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#                     run.font.size = Pt(9)
#                     if i < len(lines) - 1:
#                         run.add_break()  # 每个词性占一行
#
#             # 设置第一列宽度
#             cells[0].width = 12800
#
#         doc.save(output_filename)
#         print(f"✅ 已保存为 {os.path.abspath(output_filename)}")
#         return True
#
#     except Exception as e:
#         print(f"❌ 生成无词意文档失败：{e}")
#         return False
#
#
# def generate_dictation_books(input_file='word.txt'):
#     """生成听写本的主函数"""
#     try:
#         # 读取单词
#         words = read_words_from_file(input_file)
#         if not words:
#             return False, "没有找到单词数据"
#
#         # 获取每个单词的详细信息
#         words_with_details = []
#         for word in words:
#             meaning = get_word_details(word)
#             words_with_details.append((word, meaning))
#
#         # 生成Word文档
#         success1 = create_word_doc(words_with_details, '单词听写本（带词意）.docx')
#
#         # 生成不带词意的 Word 文档（供默写）
#         success2 = create_blank_word_doc(words_with_details, '单词听写本（无词意）.docx')
#
#         if success1 and success2:
#             return True, "听写本生成成功！已创建两个文件：单词听写本（带词意）.docx 和 单词听写本（无词意）.docx"
#         else:
#             return False, "部分文件生成失败"
#
#     except Exception as e:
#         return False, f"生成听写本失败：{str(e)}"